"""
Advanced Report Generation Module
Generates professional penetration testing reports with:
- Bootstrap 5 HTML with dark mode
- Chart.js visualizations
- CVSS-like risk scoring
- PDF export (requires wkhtmltopdf)
- DOCX export with embedded charts
"""

import json
import os
import base64
from datetime import datetime
from pathlib import Path
from jinja2 import Template
import io

try:
    import pdfkit
    HAS_PDFKIT = True
except ImportError:
    HAS_PDFKIT = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib.pyplot as plt
    import matplotlib
    matplotlib.use('Agg')
    HAS_MATPLOTLIB = True
except ImportError:
    HAS_MATPLOTLIB = False

from utils.logger import get_logger

class RiskScorer:
    """Calculate CVSS-like risk scores for findings"""
    
    SEVERITY_SCORES = {
        'critical': 9.0,
        'high': 7.5,
        'medium': 5.0,
        'low': 3.0,
        'info': 0.1
    }
    
    @staticmethod
    def calculate_finding_score(finding):
        """Calculate risk score for a single finding (0-10)"""
        severity = finding.get('severity', 'info').lower()
        base_score = RiskScorer.SEVERITY_SCORES.get(severity, 0)
        
        exploitability = finding.get('exploitability', 0.5)
        impact = finding.get('impact', 0.5)
        
        final_score = base_score * (exploitability * impact)
        return min(10.0, final_score)
    
    @staticmethod
    def calculate_overall_risk(findings):
        """Calculate overall risk score from all findings (0-10)"""
        if not findings:
            return 0.0
        
        scores = [RiskScorer.calculate_finding_score(f) for f in findings]
        avg_score = sum(scores) / len(scores)
        critical_count = sum(1 for f in findings if f.get('severity', '').lower() == 'critical')
        
        # Increase score if there are critical findings
        penalty = min(2.0, critical_count * 0.5)
        overall = min(10.0, avg_score + penalty)
        
        return round(overall, 1)
    
    @staticmethod
    def get_risk_level(score):
        """Get risk level label from score"""
        if score >= 9.0:
            return "CRITICAL"
        elif score >= 7.0:
            return "HIGH"
        elif score >= 5.0:
            return "MEDIUM"
        elif score >= 3.0:
            return "LOW"
        else:
            return "INFO"

class ChartGenerator:
    """Generate charts for reports"""
    
    @staticmethod
    def generate_port_chart(findings, output_path):
        """Generate bar chart of open ports and save as PNG"""
        if not HAS_MATPLOTLIB:
            return None
        
        port_counts = {}
        for finding in findings:
            port = finding.get('port', 'Unknown')
            service = finding.get('service', 'Unknown')
            key = f"{port}/{service}"
            port_counts[key] = port_counts.get(key, 0) + 1
        
        if not port_counts:
            return None
        
        try:
            fig, ax = plt.subplots(figsize=(12, 6), facecolor='white')
            ports = list(port_counts.keys())
            counts = list(port_counts.values())
            colors = ['#e74c3c' if c >= 3 else '#f39c12' if c >= 2 else '#27ae60' for c in counts]
            
            ax.bar(ports, counts, color=colors, edgecolor='black', linewidth=1.5)
            ax.set_xlabel('Port/Service', fontsize=12, fontweight='bold')
            ax.set_ylabel('Number of Findings', fontsize=12, fontweight='bold')
            ax.set_title('Vulnerability Distribution by Port', fontsize=14, fontweight='bold')
            ax.grid(axis='y', alpha=0.3)
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=100, bbox_inches='tight')
            plt.close()
            
            return output_path
        except Exception as e:
            print(f"Error generating chart: {e}")
            return None
    
    @staticmethod
    def generate_severity_chart(findings, output_path):
        """Generate pie chart of severity distribution"""
        if not HAS_MATPLOTLIB:
            return None
        
        severity_counts = {}
        for finding in findings:
            severity = finding.get('severity', 'info').lower()
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        if not severity_counts:
            return None
        
        try:
            colors_map = {
                'critical': '#c0392b',
                'high': '#e74c3c',
                'medium': '#f39c12',
                'low': '#f1c40f',
                'info': '#3498db'
            }
            colors = [colors_map.get(s, '#95a5a6') for s in severity_counts.keys()]
            
            fig, ax = plt.subplots(figsize=(10, 8), facecolor='white')
            ax.pie(severity_counts.values(), labels=severity_counts.keys(), autopct='%1.1f%%',
                   colors=colors, startangle=90, textprops={'fontsize': 11, 'weight': 'bold'})
            ax.set_title('Findings by Severity', fontsize=14, fontweight='bold')
            
            plt.tight_layout()
            plt.savefig(output_path, dpi=100, bbox_inches='tight')
            plt.close()
            
            return output_path
        except Exception as e:
            print(f"Error generating severity chart: {e}")
            return None

class AdvancedReportGenerator:
    """Generate professional penetration testing reports in multiple formats"""
    
    HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Penetration Test Report</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <script src="https://cdn.jsdelivr.net/npm/chart.js@4.4.0/dist/chart.umd.min.js"></script>
    <style>
        :root {
            --primary: #2c3e50;
            --success: #27ae60;
            --danger: #e74c3c;
            --warning: #f39c12;
            --info: #3498db;
            --light: #ecf0f1;
            --dark: #2c3e50;
        }
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            background: #f8f9fa;
            transition: background-color 0.3s, color 0.3s;
        }
        body.dark-mode {
            background: #1a1a1a;
            color: #e0e0e0;
        }
        .navbar-custom {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        }
        .report-header {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 40px 0;
            margin-bottom: 40px;
            text-align: center;
        }
        body.dark-mode .report-header {
            background: linear-gradient(135deg, #1a237e 0%, #4a148c 100%);
        }
        .risk-score {
            font-size: 3em;
            font-weight: bold;
            text-align: center;
        }
        .risk-critical { color: #c0392b; }
        .risk-high { color: #e74c3c; }
        .risk-medium { color: #f39c12; }
        .risk-low { color: #f1c40f; }
        .risk-info { color: #3498db; }
        .finding-card {
            border-left: 5px solid #3498db;
            margin-bottom: 20px;
            transition: box-shadow 0.3s;
        }
        .finding-card:hover {
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }
        .finding-critical { border-left-color: #c0392b; }
        .finding-high { border-left-color: #e74c3c; }
        .finding-medium { border-left-color: #f39c12; }
        .finding-low { border-left-color: #f1c40f; }
        .finding-info { border-left-color: #3498db; }
        .dark-mode .finding-card {
            background: #2a2a2a;
        }
        .summary-stat {
            text-align: center;
            padding: 20px;
            background: white;
            border-radius: 8px;
            margin-bottom: 15px;
        }
        .dark-mode .summary-stat {
            background: #2a2a2a;
        }
        .summary-stat h3 {
            color: #667eea;
            font-size: 2.5em;
            margin: 0;
        }
        .summary-stat p {
            color: #7f8c8d;
            margin-top: 5px;
        }
        .dark-mode .summary-stat p {
            color: #b0b0b0;
        }
        .table-responsive {
            border-radius: 8px;
            overflow: hidden;
        }
        .chart-container {
            position: relative;
            height: 400px;
            margin: 30px 0;
        }
        .dark-toggle {
            position: fixed;
            top: 20px;
            right: 20px;
            z-index: 1000;
        }
        @media (max-width: 768px) {
            .dark-toggle {
                position: absolute;
            }
            .risk-score {
                font-size: 2em;
            }
        }
        .footer {
            background: #f8f9fa;
            padding: 20px 0;
            margin-top: 40px;
            text-align: center;
            color: #7f8c8d;
            border-top: 1px solid #ecf0f1;
        }
        .dark-mode .footer {
            background: #2a2a2a;
            border-top-color: #404040;
            color: #b0b0b0;
        }
    </style>
</head>
<body>
    <button class="btn btn-dark dark-toggle" onclick="toggleDarkMode()">🌙 Dark Mode</button>
    
    <div class="report-header">
        <div class="container">
            <h1>Penetration Test Report</h1>
            <p class="lead">{{ timestamp }}</p>
            <p><strong>Target:</strong> {{ target }}</p>
        </div>
    </div>
    
    <div class="container">
        <!-- Risk Score Summary -->
        <div class="row mb-40">
            <div class="col-md-12">
                <h2 class="mb-3">Overall Risk Assessment</h2>
                <div class="card">
                    <div class="card-body text-center py-5">
                        <div class="risk-score risk-{{ risk_level|lower }}">{{ overall_risk }}/10</div>
                        <h3 class="mt-3">{{ risk_level }} Risk</h3>
                        <p class="text-muted">{{ total_findings }} findings identified across {{ total_hosts }} hosts</p>
                    </div>
                </div>
            </div>
        </div>
        
        <!-- Statistics -->
        <div class="row mb-40">
            <div class="col-md-3">
                <div class="summary-stat">
                    <h3>{{ total_findings }}</h3>
                    <p>Total Findings</p>
                </div>
            </div>
            <div class="col-md-3">
                <div class="summary-stat">
                    <h3>{{ critical_count }}</h3>
                    <p>Critical Issues</p>
                </div>
            </div>
            <div class="col-md-3">
                <div class="summary-stat">
                    <h3>{{ high_count }}</h3>
                    <p>High Issues</p>
                </div>
            </div>
            <div class="col-md-3">
                <div class="summary-stat">
                    <h3>{{ total_hosts }}</h3>
                    <p>Hosts Scanned</p>
                </div>
            </div>
        </div>
        
        <!-- Charts -->
        {% if port_chart_data %}
        <div class="row mb-40">
            <div class="col-md-12">
                <h2 class="mb-3">Findings Distribution</h2>
                <div class="card">
                    <div class="card-body">
                        <div class="chart-container">
                            <canvas id="portChart"></canvas>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        {% endif %}
        
        {% if severity_chart_data %}
        <div class="row mb-40">
            <div class="col-md-12">
                <h2 class="mb-3">Severity Distribution</h2>
                <div class="card">
                    <div class="card-body">
                        <div class="chart-container" style="height: 350px;">
                            <canvas id="severityChart"></canvas>
                        </div>
                    </div>
                </div>
            </div>
        </div>
        {% endif %}
        
        <!-- Findings Table -->
        <div class="row mb-40">
            <div class="col-md-12">
                <h2 class="mb-3">Detailed Findings</h2>
                <div class="table-responsive">
                    <table class="table table-hover">
                        <thead class="table-dark">
                            <tr>
                                <th>Host</th>
                                <th>Port</th>
                                <th>Service</th>
                                <th>Severity</th>
                                <th>Finding</th>
                                <th>Score</th>
                            </tr>
                        </thead>
                        <tbody>
                            {% for finding in findings %}
                            <tr>
                                <td><code>{{ finding.host }}</code></td>
                                <td>{{ finding.port }}</td>
                                <td>{{ finding.service }}</td>
                                <td><span class="badge bg-{{ severity_badge(finding.severity) }}">{{ finding.severity|upper }}</span></td>
                                <td>{{ finding.vulnerability }}</td>
                                <td><strong>{{ finding.score }}/10</strong></td>
                            </tr>
                            {% endfor %}
                        </tbody>
                    </table>
                </div>
            </div>
        </div>
        
        <!-- Individual Findings -->
        <div class="row mb-40">
            <div class="col-md-12">
                <h2 class="mb-3">Detailed Findings Analysis</h2>
                {% for finding in findings %}
                <div class="card finding-card finding-{{ finding.severity|lower }}">
                    <div class="card-header bg-light">
                        <h5 class="mb-0">{{ finding.vulnerability }}</h5>
                    </div>
                    <div class="card-body">
                        <div class="row">
                            <div class="col-md-6">
                                <p><strong>Host:</strong> <code>{{ finding.host }}</code></p>
                                <p><strong>Port:</strong> {{ finding.port }}/{{ finding.service }}</p>
                                <p><strong>Severity:</strong> <span class="badge bg-{{ severity_badge(finding.severity) }}">{{ finding.severity|upper }}</span></p>
                            </div>
                            <div class="col-md-6">
                                <p><strong>Risk Score:</strong> <strong class="text-danger">{{ finding.score }}/10</strong></p>
                                <p><strong>Exploitability:</strong> {{ (finding.exploitability * 100)|int }}%</p>
                                <p><strong>Impact:</strong> {{ (finding.impact * 100)|int }}%</p>
                            </div>
                        </div>
                        <hr>
                        <p><strong>Description:</strong></p>
                        <p>{{ finding.note }}</p>
                    </div>
                </div>
                {% endfor %}
            </div>
        </div>
        
        <!-- Recommendations -->
        <div class="row mb-40">
            <div class="col-md-12">
                <h2 class="mb-3">Recommendations</h2>
                <div class="card">
                    <div class="card-body">
                        <ul>
                            <li><strong>Address Critical Issues:</strong> Immediately patch all critical-severity vulnerabilities</li>
                            <li><strong>Network Segmentation:</strong> Implement network segmentation to isolate sensitive systems</li>
                            <li><strong>SMB Hardening:</strong> Disable SMBv1, enable SMB signing, and restrict shares</li>
                            <li><strong>Patch Management:</strong> Establish regular patching schedule for OS and applications</li>
                            <li><strong>Access Control:</strong> Review and enforce principle of least privilege (PoLP)</li>
                            <li><strong>Monitoring:</strong> Implement EDR and SIEM for continuous monitoring</li>
                            <li><strong>Security Awareness:</strong> Conduct security training for all staff members</li>
                        </ul>
                    </div>
                </div>
            </div>
        </div>
    </div>
    
    <div class="footer">
        <p>&copy; {{ copyright_year }} VulnScan Pentest Pro - Professional Security Assessment</p>
        <p>This report is confidential and intended for authorized personnel only.</p>
    </div>
    
    <script>
        function toggleDarkMode() {
            document.body.classList.toggle('dark-mode');
            localStorage.setItem('darkMode', document.body.classList.contains('dark-mode'));
        }
        
        if (localStorage.getItem('darkMode') === 'true') {
            document.body.classList.add('dark-mode');
        }
        
        // Port Chart
        {% if port_chart_data %}
        const portCtx = document.getElementById('portChart').getContext('2d');
        new Chart(portCtx, {
            type: 'bar',
            data: {{ port_chart_data|safe }},
            options: {
                responsive: true,
                maintainAspectRatio: false,
                plugins: {
                    legend: { display: false }
                },
                scales: {
                    y: { beginAtZero: true }
                }
            }
        });
        {% endif %}
        
        // Severity Chart
        {% if severity_chart_data %}
        const severityCtx = document.getElementById('severityChart').getContext('2d');
        new Chart(severityCtx, {
            type: 'doughnut',
            data: {{ severity_chart_data|safe }},
            options: {
                responsive: true,
                maintainAspectRatio: false,
                plugins: {
                    legend: { position: 'bottom' }
                }
            }
        });
        {% endif %}
    </script>
</body>
</html>"""
    
    def __init__(self, output_dir):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.logger = get_logger(__name__)
    
    def generate_reports(self, findings, basename):
        """Generate all report formats and return file paths"""
        if not findings:
            self.logger.warning("No findings provided for report generation")
            return {}
        
        results = {}
        
        html_file = self._generate_html_report(findings, basename)
        if html_file:
            results['html'] = str(html_file)
        
        if HAS_PDFKIT and html_file:
            pdf_file = self._generate_pdf_report(html_file, basename)
            if pdf_file:
                results['pdf'] = str(pdf_file)
        
        if HAS_DOCX:
            docx_file = self._generate_docx_report(findings, basename)
            if docx_file:
                results['docx'] = str(docx_file)
        
        self.logger.info(f"Reports generated: {', '.join(results.keys())}")
        return results
    
    def _generate_html_report(self, findings, basename):
        """Generate HTML report with Bootstrap and Chart.js"""
        try:
            # Calculate statistics
            total_findings = len(findings)
            total_hosts = len(set(f.get('host', 'Unknown') for f in findings))
            critical_count = sum(1 for f in findings if f.get('severity', '').lower() == 'critical')
            high_count = sum(1 for f in findings if f.get('severity', '').lower() == 'high')
            
            # Calculate overall risk score
            overall_risk = RiskScorer.calculate_overall_risk(findings)
            risk_level = RiskScorer.get_risk_level(overall_risk)
            
            # Add individual scores to findings
            for finding in findings:
                finding['score'] = round(RiskScorer.calculate_finding_score(finding), 1)
            
            # Prepare chart data
            port_chart_data = self._prepare_port_chart_data(findings)
            severity_chart_data = self._prepare_severity_chart_data(findings)
            
            # Render template
            template = Template(self.HTML_TEMPLATE)
            html_content = template.render(
                timestamp=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                target=', '.join(set(f.get('host', 'Unknown') for f in findings)),
                total_findings=total_findings,
                total_hosts=total_hosts,
                critical_count=critical_count,
                high_count=high_count,
                overall_risk=overall_risk,
                risk_level=risk_level,
                findings=findings,
                port_chart_data=port_chart_data,
                severity_chart_data=severity_chart_data,
                copyright_year=datetime.now().year,
                severity_badge=lambda sev: {'critical': 'danger', 'high': 'danger', 
                                           'medium': 'warning', 'low': 'warning', 
                                           'info': 'info'}.get(sev.lower(), 'secondary')
            )
            
            html_file = self.output_dir / f"{basename}.html"
            with open(html_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            self.logger.info(f"HTML report generated: {html_file}")
            return html_file
            
        except Exception as e:
            self.logger.error(f"Error generating HTML report: {e}")
            return None
    
    def _prepare_port_chart_data(self, findings):
        """Prepare data for port chart.js visualization"""
        port_counts = {}
        for finding in findings:
            port = finding.get('port', 'Unknown')
            service = finding.get('service', 'Unknown')
            key = f"{port}/{service}"
            port_counts[key] = port_counts.get(key, 0) + 1
        
        labels = list(port_counts.keys())
        data = list(port_counts.values())
        
        return json.dumps({
            'labels': labels,
            'datasets': [{
                'label': 'Number of Findings',
                'data': data,
                'backgroundColor': ['#e74c3c', '#f39c12', '#27ae60'],
                'borderColor': '#2c3e50',
                'borderWidth': 2
            }]
        })
    
    def _prepare_severity_chart_data(self, findings):
        """Prepare data for severity pie chart"""
        severity_counts = {}
        for finding in findings:
            severity = finding.get('severity', 'info').lower()
            severity_counts[severity] = severity_counts.get(severity, 0) + 1
        
        color_map = {
            'critical': '#c0392b',
            'high': '#e74c3c',
            'medium': '#f39c12',
            'low': '#f1c40f',
            'info': '#3498db'
        }
        
        return json.dumps({
            'labels': list(severity_counts.keys()),
            'datasets': [{
                'data': list(severity_counts.values()),
                'backgroundColor': [color_map.get(k, '#95a5a6') for k in severity_counts.keys()],
                'borderColor': '#fff',
                'borderWidth': 2
            }]
        })
    
    def _generate_pdf_report(self, html_file, basename):
        """Generate PDF from HTML using pdfkit (requires wkhtmltopdf)"""
        if not HAS_PDFKIT:
            self.logger.warning("pdfkit not installed, skipping PDF generation")
            return None
        
        try:
            pdf_file = self.output_dir / f"{basename}.pdf"
            options = {
                'javascript-delay': '1000',
                'enable-local-file-access': None
            }
            pdfkit.from_file(str(html_file), str(pdf_file), options=options)
            self.logger.info(f"PDF report generated: {pdf_file}")
            return pdf_file
        except Exception as e:
            self.logger.error(f"Error generating PDF report: {e}. Ensure wkhtmltopdf is installed.")
            return None
    
    def _generate_docx_report(self, findings, basename):
        """Generate DOCX report with embedded charts"""
        if not HAS_DOCX:
            self.logger.warning("python-docx not installed, skipping DOCX generation")
            return None
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Penetration Test Report', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"Targets: {', '.join(set(f.get('host', 'Unknown') for f in findings))}")
            
            # Summary
            overall_risk = RiskScorer.calculate_overall_risk(findings)
            risk_level = RiskScorer.get_risk_level(overall_risk)
            doc.add_heading('Executive Summary', level=1)
            summary_table = doc.add_table(rows=5, cols=2)
            summary_table.rows[0].cells[0].text = 'Metric'
            summary_table.rows[0].cells[1].text = 'Value'
            summary_table.rows[1].cells[0].text = 'Overall Risk Score'
            summary_table.rows[1].cells[1].text = f'{overall_risk}/10 ({risk_level})'
            summary_table.rows[2].cells[0].text = 'Total Findings'
            summary_table.rows[2].cells[1].text = str(len(findings))
            summary_table.rows[3].cells[0].text = 'Critical Issues'
            summary_table.rows[3].cells[1].text = str(sum(1 for f in findings if f.get('severity', '').lower() == 'critical'))
            summary_table.rows[4].cells[0].text = 'Hosts Scanned'
            summary_table.rows[4].cells[1].text = str(len(set(f.get('host', 'Unknown') for f in findings)))
            
            # Findings Table
            doc.add_heading('Findings', level=1)
            findings_table = doc.add_table(rows=len(findings) + 1, cols=5)
            findings_table.style = 'Light Grid Accent 1'
            header_cells = findings_table.rows[0].cells
            header_cells[0].text = 'Host'
            header_cells[1].text = 'Port'
            header_cells[2].text = 'Service'
            header_cells[3].text = 'Severity'
            header_cells[4].text = 'Finding'
            
            for i, finding in enumerate(findings, start=1):
                row_cells = findings_table.rows[i].cells
                row_cells[0].text = finding.get('host', 'Unknown')
                row_cells[1].text = str(finding.get('port', 'N/A'))
                row_cells[2].text = finding.get('service', 'Unknown')
                row_cells[3].text = finding.get('severity', 'info').upper()
                row_cells[4].text = finding.get('vulnerability', 'Unknown')
            
            # Recommendations
            doc.add_heading('Recommendations', level=1)
            doc.add_paragraph('Address critical issues immediately', style='List Bullet')
            doc.add_paragraph('Implement network segmentation', style='List Bullet')
            doc.add_paragraph('Establish regular patching schedule', style='List Bullet')
            doc.add_paragraph('Enable enhanced monitoring and logging', style='List Bullet')
            
            docx_file = self.output_dir / f"{basename}.docx"
            doc.save(str(docx_file))
            self.logger.info(f"DOCX report generated: {docx_file}")
            return docx_file
            
        except Exception as e:
            self.logger.error(f"Error generating DOCX report: {e}")
            return None


def generate_reports(findings, basename):
    """Convenience function to generate all report formats"""
    output_dir = str(Path(basename).parent)
    basename_only = Path(basename).name
    
    generator = AdvancedReportGenerator(output_dir)
    return generator.generate_reports(findings, basename_only)
